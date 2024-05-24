import sys
import json
import os
import difflib
from docx import Document
from io import BytesIO

def read_docx(file_path):
    try:
        # Mở tệp .docx
        doc = Document(file_path)

        # Lấy nội dung từ các đoạn văn bản
        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)

        return "\n".join(content)
    except Exception as e:
        print("Lỗi:", e)
        return None
    
def compare_files(file1_lines, file2_lines):
  matcher = difflib.SequenceMatcher(None, file1_lines, file2_lines)
  similarity_ratio = matcher.ratio() * 100  # Tỉ lệ phần trăm độ tương đồng
  return similarity_ratio

json_args = sys.argv[1]
listUploadFile = json.loads(json_args)
contentUploads = []

for file in listUploadFile:
    if file["mimetype"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        buffer = file["buffer"]
        data = buffer["data"]
        file_content = bytes(data)
        file_stream = BytesIO(file_content)
        doc = Document(file_stream)
        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
        newContent = "\n".join(content)
        newObj = {
            "name": file["originalname"],
            "content": newContent
        }
        contentUploads.append(newObj)

current_file_path = os.path.abspath(__file__)
folderPath = os.path.normpath(os.path.join(os.path.dirname(current_file_path),'..','Documents'))
foldersName = os.listdir(folderPath)
kq = []
for i in range(1,len(foldersName)):
    if len(contentUploads) == 0:
        break;
    else:
        folderPath = os.path.normpath(os.path.join(os.path.dirname(current_file_path),'..','Documents',foldersName[i]))
        listFile = os.listdir(folderPath)
        for file in listFile:
            path = os.path.normpath(os.path.join(os.path.dirname(current_file_path),'..','Documents',foldersName[i],file))
            if os.path.splitext(path)[1].lower() == '.docx':
                content = read_docx(path)
                for uploadFile in contentUploads:
                    similarity = compare_files(content, uploadFile["content"])
                    if similarity > 60:
                        kq.append(uploadFile["name"])
                        contentUploads   = [x for x in contentUploads if x["name"] != uploadFile["name"]]

result = json.dumps(kq)
print(result)